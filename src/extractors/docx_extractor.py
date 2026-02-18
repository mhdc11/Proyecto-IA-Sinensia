"""
DOCX Text Extractor - Analizador de Documentos Legales

Extractor de texto para archivos Microsoft Word DOCX usando python-docx.

Author: Analizador de Documentos Legales Team
Date: 2026-02-18
"""

from pathlib import Path
from typing import Optional

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not available. Install with: pip install python-docx")


def extract_text_docx(file_path: Path) -> str:
    """
    Extrae texto de un archivo DOCX usando python-docx

    Args:
        file_path: Ruta al archivo DOCX

    Returns:
        str: Texto extraído del documento

    Raises:
        FileNotFoundError: Si el archivo no existe
        ValueError: Si el archivo no es un DOCX válido
        RuntimeError: Si python-docx no está instalado o falla la extracción

    Example:
        >>> from pathlib import Path
        >>> texto = extract_text_docx(Path("contrato.docx"))
        >>> print(f"Extracted {len(texto)} characters")
        Extracted 15620 characters
    """
    if not file_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    if file_path.suffix.lower() not in [".docx", ".doc"]:
        raise ValueError(f"File is not a DOCX/DOC: {file_path}")

    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx not installed. Install with: pip install python-docx"
        )

    try:
        doc = DocxDocument(str(file_path))

        # Extraer texto de todos los párrafos
        paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Solo añadir párrafos no vacíos
                paragraphs.append(text)

        # Extraer texto de tablas (común en documentos legales)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        # Unir párrafos con saltos de línea dobles para preservar estructura
        full_text = "\n\n".join(paragraphs)

        if not full_text.strip():
            raise RuntimeError(f"No text extracted from DOCX. File may be empty.")

        return full_text

    except Exception as e:
        if "python-docx" in str(e) or "docx" in str(e).lower():
            raise RuntimeError(f"Error extracting text from DOCX: {e}") from e
        raise ValueError(f"Invalid or corrupted DOCX file: {e}") from e


def extract_text_with_metadata(file_path: Path) -> dict:
    """
    Extrae texto y metadata de un DOCX

    Args:
        file_path: Ruta al archivo DOCX

    Returns:
        dict: Diccionario con:
            - text: Texto extraído
            - paragraphs_count: Número de párrafos
            - tables_count: Número de tablas
            - core_properties: Metadata del documento (autor, título, etc.)

    Example:
        >>> result = extract_text_with_metadata(Path("contrato.docx"))
        >>> print(result['paragraphs_count'])
        45
        >>> print(result['core_properties']['author'])
        'Juan Pérez'
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")

    doc = DocxDocument(str(file_path))

    # Extraer texto
    text = extract_text_docx(file_path)

    # Contar elementos
    paragraphs_count = len([p for p in doc.paragraphs if p.text.strip()])
    tables_count = len(doc.tables)

    # Extraer core properties (metadata)
    core_props = {}
    if hasattr(doc, "core_properties"):
        props = doc.core_properties
        core_props = {
            "author": props.author,
            "title": props.title,
            "subject": props.subject,
            "created": props.created.isoformat() if props.created else None,
            "modified": props.modified.isoformat() if props.modified else None,
        }

    return {
        "text": text,
        "paragraphs_count": paragraphs_count,
        "tables_count": tables_count,
        "core_properties": core_props,
    }


if __name__ == "__main__":
    # Test de extracción de DOCX
    print("=" * 60)
    print("Testing DOCX Text Extraction")
    print("=" * 60)

    print("\n📋 Module loaded successfully")
    print(f"✅ python-docx available: {DOCX_AVAILABLE}")

    print("\n📋 Example usage:")
    print("""
    from pathlib import Path
    from src.extractors.docx_extractor import extract_text_docx

    # Extraer texto simple
    texto = extract_text_docx(Path("contrato.docx"))
    print(f"Extracted {len(texto)} characters")

    # Extraer con metadata
    result = extract_text_with_metadata(Path("contrato.docx"))
    print(f"Paragraphs: {result['paragraphs_count']}")
    print(f"Tables: {result['tables_count']}")
    print(f"Author: {result['core_properties']['author']}")
    """)

    # Test de validación básica
    print("\n📋 Testing error handling:")
    try:
        extract_text_docx(Path("nonexistent.docx"))
        print("❌ Should have raised FileNotFoundError")
    except FileNotFoundError:
        print("✅ Correctly raised FileNotFoundError")

    try:
        extract_text_docx(Path("not_a_docx.txt"))
        print("❌ Should have raised ValueError")
    except ValueError:
        print("✅ Correctly raised ValueError for non-DOCX file")

    print("\n✅ DOCX extractor ready!")
    print("⚠️  Note: Real testing requires actual DOCX files in tests/fixtures/")
